from docx import Document

def convert_docx_to_txt(docx_file, output_file):
    try:
        
        doc = Document(docx_file)
        
        
        with open(output_file, 'w', encoding='utf-8') as txt_file:
            paragraph_count = 0
            table_count = 0
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if text: 
                    txt_file.write(text + '\n')
                    paragraph_count += 1

                if paragraph_count % 100 == 0:
                    print(f"Paragraph {paragraph_count} elaborated.")
                    
            
            txt_file.write('\n')
            for table in doc.tables:
                table_count += 1

                txt_file.write('\n')

                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            cell_text = para.text.strip()
                            if not cell_text:  
                                cell_text = "no"
                            row_text.append(cell_text)
                    
                    if row_text:
                        txt_file.write(" : ".join(row_text) + '\n')
                
                txt_file.write('\n')

                print(f"Table {table_count} elaborated.")
        
        print(f"Conversion completed: {output_file}")
        print(f"Total paragraphs: {paragraph_count}")
        print(f"Total tables: {table_count}")
    
    except Exception as e:
        print(f"Error: {e}")
